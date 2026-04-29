"""Document parser — extracts control statements from policy documents."""

import re
from pathlib import Path
from typing import Optional

from .models import ParsedDocument, ControlStatement, CoverageStatus


def parse_policy_document(filepath: str) -> ParsedDocument:
    """Parse a policy document (.txt, .md, .pdf, .docx) and extract control statements.

    For .pdf and .docx, optional dependencies are required:
      - pdfplumber or PyPDF2 for PDF
      - python-docx for DOCX

    Plain text (.txt, .md) files are read directly with no extra dependencies.
    """
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    suffix = path.suffix.lower()

    if suffix in (".txt", ".md", ""):
        text = path.read_text(encoding="utf-8", errors="replace")
    elif suffix == ".pdf":
        text = _read_pdf(filepath)
    elif suffix == ".docx":
        text = _read_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: .txt, .md, .pdf, .docx")

    paragraphs = _split_paragraphs(text)
    extracted = _extract_controls(paragraphs)

    return ParsedDocument(
        source_file=str(path.absolute()),
        paragraphs=paragraphs,
        extracted_controls=extracted,
    )


def _split_paragraphs(text: str) -> list[str]:
    """Split text into non-empty paragraphs."""
    paragraphs = []
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped:
            paragraphs.append(stripped)
    return paragraphs


# Regex patterns for identifying control-like statements
_CONTROL_PATTERNS = [
    re.compile(r"(?:Control|Requirement|Policy)[\s:]+(.+)", re.IGNORECASE),
    re.compile(r"(?:shall|must|will)\s+(.+)", re.IGNORECASE),
    re.compile(r"^(?:[A-Z]\.\d+\.\d+[\s:]+)(.+)", re.MULTILINE),       # ISO-style: A.9.2.1
    re.compile(r"^(?:[A-Z]{2}\.[A-Z]{2}-\d+[\s:]+)(.+)", re.MULTILINE), # NIST: ID.AM-2
    re.compile(r"^(?:CIS\s*\d+\.\d+[\s:]+)(.+)", re.MULTILINE),         # CIS: CIS 1.1
    re.compile(r"^(?:VOD-\w{3}-\d+[\s:]+)(.+)", re.MULTILINE),          # VODAFONE: VOD-ACC-001
]


def _extract_controls(paragraphs: list[str]) -> list[ControlStatement]:
    """Extract control-like statements from paragraphs.

    Uses regex patterns to identify sentences that look like security controls
    (containing keywords like "shall", "must", "Control:", etc.).
    """
    extracted = []
    seen = set()

    for i, para in enumerate(paragraphs):
        for pattern in _CONTROL_PATTERNS:
            match = pattern.search(para)
            if match:
                desc = para.strip()
                # De-duplicate near-identical paragraphs
                key = desc[:80].lower()
                if key in seen:
                    continue
                seen.add(key)

                extracted.append(ControlStatement(
                    framework="PARSED",
                    control_id=f"PARSED-{i + 1:03d}",
                    title=desc[:100] + ("..." if len(desc) > 100 else ""),
                    description=desc,
                    category="Extracted",
                    status=CoverageStatus.GAP,
                ))
                break

    return extracted


def _read_pdf(filepath: str) -> str:
    """Extract text from a PDF file using pdfplumber or PyPDF2."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        raise ImportError(
            "PDF support requires 'pdfplumber' or 'PyPDF2'. "
            "Install with: pip install pdfplumber"
        )


def _read_docx(filepath: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text += "\n" + cell.text.strip()

        return full_text
    except ImportError:
        raise ImportError(
            "DOCX support requires 'python-docx'. "
            "Install with: pip install python-docx"
        )
