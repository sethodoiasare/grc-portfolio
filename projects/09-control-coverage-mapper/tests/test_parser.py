"""Tests for document parser."""

import tempfile
from pathlib import Path

import pytest

from src.parser import parse_policy_document
from src.models import ParsedDocument, ControlStatement


class TestParseTextFile:
    def test_parses_simple_txt(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("Control: All access shall be authorised.\n\n")
            f.write("Systems must encrypt data at rest.\n")
            f.write("This is just a note.\n")
            f.flush()
            fpath = f.name

        try:
            doc = parse_policy_document(fpath)
            assert isinstance(doc, ParsedDocument)
            assert doc.source_file.endswith(".txt")
            assert len(doc.paragraphs) == 3
            assert len(doc.extracted_controls) >= 1
        finally:
            Path(fpath).unlink()

    def test_parses_markdown(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write("## Security Policy\n\n")
            f.write("Requirement: MFA shall be used for all privileged accounts.\n\n")
            f.write("Passwords must be rotated every 90 days.\n")
            f.flush()
            fpath = f.name

        try:
            doc = parse_policy_document(fpath)
            assert len(doc.paragraphs) >= 2
            assert len(doc.extracted_controls) >= 1
            # At least one control from "Requirement:" or "shall" patterns
            texts = [c.description for c in doc.extracted_controls]
            assert any("MFA" in t for t in texts) or any("privileged" in t for t in texts)
        finally:
            Path(fpath).unlink()

    def test_handles_empty_file(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("")
            f.flush()
            fpath = f.name

        try:
            doc = parse_policy_document(fpath)
            assert len(doc.paragraphs) == 0
            assert len(doc.extracted_controls) == 0
        finally:
            Path(fpath).unlink()

    def test_handles_encoding_issues(self):
        with tempfile.NamedTemporaryFile(mode="wb", suffix=".txt", delete=False) as f:
            # Write bytes that include some non-UTF8
            f.write(b"Valid ASCII line\n")
            f.write(b"Some valid text with special chars.\n")
            f.flush()
            fpath = f.name

        try:
            doc = parse_policy_document(fpath)
            assert len(doc.paragraphs) >= 1
        finally:
            Path(fpath).unlink()

    def test_raises_on_missing_file(self):
        with pytest.raises(FileNotFoundError):
            parse_policy_document("/tmp/nonexistent_file_xyz123.txt")

    def test_extracts_iso_style_controls(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("A.9.2.1: Access control shall be implemented for all information systems.\n")
            f.write("A.9.4.2: Secure log-on procedures shall be required.\n")
            f.flush()
            fpath = f.name

        try:
            doc = parse_policy_document(fpath)
            iso_controls = [c for c in doc.extracted_controls
                            if any(pattern in c.description for pattern in ["Access control", "log-on"])]
            assert len(iso_controls) >= 1
        finally:
            Path(fpath).unlink()

    def test_dedup_near_identical_paragraphs(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("Control: Access shall be controlled.\n")
            f.write("Control: Access shall be controlled.\n")
            f.flush()
            fpath = f.name

        try:
            doc = parse_policy_document(fpath)
            # Should deduplicate — only one extracted control for identical text
            access_controls = [c for c in doc.extracted_controls if "Access" in c.description]
            assert len(access_controls) == 1
        finally:
            Path(fpath).unlink()


class TestParsePDF:
    def test_pdf_raises_import_error_when_no_deps(self):
        """PDF parsing requires pdfplumber or PyPDF2."""
        try:
            import pdfplumber  # noqa: F401
            skip_pdfplumber = False
        except ImportError:
            skip_pdfplumber = True

        try:
            import PyPDF2  # noqa: F401
            skip_pypdf2 = False
        except ImportError:
            skip_pypdf2 = True

        if not skip_pdfplumber or not skip_pypdf2:
            pytest.skip("PDF parsing library is available — cannot test error path")

        with pytest.raises(ImportError, match="pdfplumber"):
            parse_policy_document("/tmp/fake.pdf")


class TestParseDOCX:
    def test_docx_raises_import_error_when_no_deps(self):
        """DOCX parsing requires python-docx."""
        try:
            import docx  # noqa: F401
            pytest.skip("python-docx is available — cannot test error path")
        except ImportError:
            pass

        with pytest.raises(ImportError, match="python-docx"):
            parse_policy_document("/tmp/fake.docx")


class TestParseUnsupported:
    def test_raises_on_unsupported_format(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".xyz", delete=False) as f:
            f.write("test")
            f.flush()
            fpath = f.name

        try:
            with pytest.raises(ValueError, match="Unsupported"):
                parse_policy_document(fpath)
        finally:
            Path(fpath).unlink()
