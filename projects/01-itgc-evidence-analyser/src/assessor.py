"""
Evidence Assessor

Orchestrates the full ITGC evidence assessment pipeline:

  1. Validate the requested control exists in the controls dataset.
  2. Format the control definition as a prompt-ready string.
  3. Dispatch to GRCClaudeClient for AI-powered assessment.
  4. Parse the raw Claude response dict into a strongly-typed AssessmentResult.

Supported evidence ingestion formats
-------------------------------------
- Plain text / Markdown  (.txt, .md)
- PDF                    (.pdf)  — text extracted via pdfplumber
- CSV                    (.csv)  — summarised via pandas

All public methods raise ValueError for unknown control IDs rather than
silently returning None, so the caller can surface a clear error message.
"""

import pdfplumber
import pandas as pd
from pathlib import Path
from datetime import datetime

from src.control_parser import ControlParser
from src.claude_client import GRCClaudeClient, SONNET
from src.models import (
    AssessmentResult,
    Verdict,
    RiskRating,
    StatementType,
    DraftFinding,
)


class EvidenceAssessor:
    """
    High-level orchestrator for evidence assessment.

    A single instance can be reused across many assessments; it holds
    lightweight, stateless helpers (ControlParser, GRCClaudeClient).
    """

    def __init__(self):
        self.parser = ControlParser()
        self.client = GRCClaudeClient()

    # ------------------------------------------------------------------
    # Primary assessment methods
    # ------------------------------------------------------------------

    def assess(
        self,
        control_id: str,
        evidence_text: str,
        statement_type: str = "D",
        target_statements: list[str] | None = None,
    ) -> AssessmentResult:
        """
        Assess a block of evidence text against a single Vodafone control.

        Parameters
        ----------
        control_id : str
            Vodafone control identifier, e.g. "IAM_001".
        evidence_text : str
            The full body of audit evidence to be evaluated.
        statement_type : str
            "D" to assess against design statements (default), or
            "E" to assess against evidence statements.
        target_statements : list[str] | None
            Optional list of specific D/E statement IDs to assess against
            (e.g., ["D1", "E1"]). When provided, only those statements are
            included in the prompt. When empty/None, all statements are used.

        Returns
        -------
        AssessmentResult
            Fully typed result including verdict, confidence, gaps, risk rating,
            optional draft finding, and token-usage metadata.

        Raises
        ------
        ValueError
            If control_id is not present in the controls dataset.
        """
        control = self.parser.get_control(control_id)
        if not control:
            available = ", ".join(c["control_id"] for c in self.parser.list_controls())
            raise ValueError(
                f"Control '{control_id}' not found. "
                f"Available controls: {available}"
            )

        control_context = self.parser.format_for_prompt(
            control_id, statement_type, target_statements=target_statements
        )
        result_dict, tokens_used = self.client.assess_evidence(
            control_context, evidence_text, statement_type
        )

        return self._build_result(control, result_dict, statement_type, tokens_used)

    def assess_batch(self, items: list[dict]) -> list[AssessmentResult]:
        """
        Assess multiple control/evidence pairs in sequence.

        Parameters
        ----------
        items : list[dict]
            Each dict must contain:
              - ``control_id``        (str)  — required
              - ``evidence_text``     (str)  — required
              - ``statement_type``    (str)  — optional, defaults to "D"
              - ``target_statements`` (list) — optional, specific D/E IDs to target

        Returns
        -------
        list[AssessmentResult]
            Results in the same order as the input list.

        Notes
        -----
        Assessment stops and raises ValueError immediately if an unknown
        control_id is encountered. Wrap the call in a try/except if partial
        results on error are preferred.
        """
        results: list[AssessmentResult] = []
        for item in items:
            result = self.assess(
                control_id=item["control_id"],
                evidence_text=item["evidence_text"],
                statement_type=item.get("statement_type", "D"),
                target_statements=item.get("target_statements") or [],
            )
            results.append(result)
        return results

    def extract_file_text(self, file_path: str) -> str:
        """
        Extract text from any supported file type. Public method usable by
        API endpoints that need extraction without running an assessment.

        Returns the extracted text as a string, or a placeholder message
        if extraction fails.
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        try:
            if suffix == ".pdf":
                return self._extract_pdf_text(path)
            elif suffix == ".csv":
                return self._extract_csv_summary(path)
            elif suffix == ".xlsx":
                return self._extract_xlsx_summary(path)
            elif suffix in (".docx", ".doc"):
                return self._extract_docx_text(path)
            elif suffix in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff"):
                return self._extract_image_metadata(path)
            else:
                try:
                    return path.read_text(encoding="utf-8")
                except UnicodeDecodeError:
                    return f"[Binary file: {path.name} ({path.stat().st_size} bytes)]"
        except Exception as e:
            return f"[Extraction failed for {path.name}: {e}]"

    def assess_from_file(
        self,
        control_id: str,
        file_path: str,
        statement_type: str = "D",
        target_statements: str = "",
    ) -> AssessmentResult:
        """
        Read evidence from a file on disk and assess it against a control.

        Supported formats
        -----------------
        - .pdf  : text extracted from all pages via pdfplumber
        - .csv  : summarised as a text block via pandas (up to 50 rows shown)
        - .xlsx : summarised from all sheets via openpyxl
        - .docx : paragraph text via python-docx
        - .png / .jpg : metadata via Pillow
        - .txt / .md / any other extension : read as UTF-8 plain text

        Parameters
        ----------
        control_id : str
            Vodafone control identifier.
        file_path : str
            Absolute or relative path to the evidence file.
        statement_type : str
            "D" (design) or "E" (evidence) — defaults to "D".
        target_statements : str
            Comma-separated D/E statement IDs, e.g. "D1,E1". When empty, all
            statements are assessed.

        Returns
        -------
        AssessmentResult

        Raises
        ------
        FileNotFoundError
            If the path does not exist.
        ValueError
            If the control_id is not recognised.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Evidence file not found: {file_path}")

        evidence_text = self.extract_file_text(file_path)
        target_list = [t.strip() for t in target_statements.split(",") if t.strip()] if target_statements else []
        return self.assess(control_id, evidence_text, statement_type, target_statements=target_list)

    # ------------------------------------------------------------------
    # File extraction helpers
    # ------------------------------------------------------------------

    def _extract_pdf_text(self, path: Path) -> str:
        """
        Extract comprehensive content from a PDF using pdfplumber, preserving
        text, tables, and page structure with metadata.
        """
        parts = [f"PDF Evidence: {path.name}"]
        with pdfplumber.open(path) as pdf:
            # PDF metadata
            if pdf.metadata:
                meta = pdf.metadata
                if meta.get("Title"):
                    parts.append(f"Title: {meta['Title']}")
                if meta.get("Author"):
                    parts.append(f"Author: {meta['Author']}")
                if meta.get("CreationDate"):
                    parts.append(f"Creation Date: {meta['CreationDate']}")
                if meta.get("ModDate"):
                    parts.append(f"Modification Date: {meta['ModDate']}")
                if meta.get("Producer"):
                    parts.append(f"Producer: {meta['Producer']}")
                parts.append("")

            total_pages = len(pdf.pages)
            parts.append(f"Total Pages: {total_pages}")
            pages_with_tables = 0
            pages_with_text = 0

            for pi, page in enumerate(pdf.pages):
                page_text_parts = []
                page_num = pi + 1

                # Tables first (they're often the most important evidence)
                tables = page.extract_tables()
                if tables:
                    pages_with_tables += 1
                    page_text_parts.append(f"\n[Page {page_num} Tables]")
                    for ti, table in enumerate(tables):
                        if table:
                            page_text_parts.append(f"Table {ti + 1}:")
                            for row in table[:60]:  # limit table rows
                                row_text = " | ".join(str(c) if c else "" for c in row)
                                page_text_parts.append(f"  {row_text}")
                            if len(table) > 60:
                                page_text_parts.append(f"  ... ({len(table) - 60} more rows)")

                # Text
                text = (page.extract_text() or "").strip()
                if text:
                    pages_with_text += 1
                    page_text_parts.append(f"\n[Page {page_num} Text]")
                    page_text_parts.append(text)

                if page_text_parts:
                    parts.append("".join(page_text_parts))

            parts.insert(3, f"Pages with text: {pages_with_text}, Pages with tables: {pages_with_tables}")

        non_empty = [p for p in parts if p.strip()]
        return "\n".join(non_empty) if non_empty else "[PDF contained no extractable text]"

    def _extract_csv_summary(self, path: Path) -> str:
        """
        Summarise a CSV file as plain text for evidence assessment.

        Includes the filename, row/column counts, column names, and up to
        50 rows of data so the model can reason about the contents without
        exceeding context limits on large files.
        """
        df = pd.read_csv(path)
        lines = [
            f"CSV Evidence Summary: {path.name}",
            f"Rows: {len(df)}, Columns: {list(df.columns)}",
            "",
            df.to_string(max_rows=50, max_cols=20),
        ]
        return "\n".join(lines)

    def _extract_xlsx_summary(self, path: Path) -> str:
        """
        Summarise an XLSX/Excel file for evidence assessment.

        Reads all sheets and includes sheet names, row/column counts,
        and up to 50 rows from each sheet.
        """
        import openpyxl

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        parts = [f"XLSX Evidence Summary: {path.name}"]
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(max_row=51, values_only=True))
            if not rows:
                parts.append(f"\nSheet '{sheet_name}': (empty)")
                continue
            header = [str(c) if c else "" for c in rows[0]]
            parts.append(f"\nSheet '{sheet_name}': {len(rows) - 1} data rows, columns: {header}")
            for row in rows[1:]:
                parts.append("  " + " | ".join(str(c) if c else "" for c in row))
        wb.close()
        return "\n".join(parts)

    def _extract_docx_text(self, path: Path) -> str:
        """
        Extract comprehensive content from a Word (.docx) document including
        paragraphs, tables, headers, footers, and embedded image metadata.
        """
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        doc = Document(str(path))
        parts = [f"DOCX Evidence: {path.name}"]
        total_parts = 0

        # Core properties
        if doc.core_properties:
            props = doc.core_properties
            meta = []
            if props.title:
                meta.append(f"Title: {props.title}")
            if props.author:
                meta.append(f"Author: {props.author}")
            if props.created:
                meta.append(f"Created: {props.created.isoformat()}")
            if props.modified:
                meta.append(f"Modified: {props.modified.isoformat()}")
            if props.last_modified_by:
                meta.append(f"Last Modified By: {props.last_modified_by}")
            if meta:
                parts.append("--- Document Properties ---")
                parts.extend(meta)
                parts.append("")

        # Headers
        header_texts = []
        for section in doc.sections:
            h = section.header
            if h and h.paragraphs:
                ht = "\n".join(p.text for p in h.paragraphs if p.text.strip())
                if ht:
                    header_texts.append(ht)
        if header_texts:
            parts.append("--- Headers ---")
            parts.extend(header_texts)
            parts.append("")

        # Footers
        footer_texts = []
        for section in doc.sections:
            f = section.footer
            if f and f.paragraphs:
                ft = "\n".join(p.text for p in f.paragraphs if p.text.strip())
                if ft:
                    footer_texts.append(ft)
        if footer_texts:
            parts.append("--- Footers ---")
            parts.extend(footer_texts)
            parts.append("")

        # Paragraphs with style info
        parts.append("--- Body Content ---")
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            style = p.style.name if p.style else "Normal"
            prefix = f"[{style}] " if style and style != "Normal" else ""
            parts.append(f"{prefix}{text}")
            total_parts += 1

        # Tables
        if doc.tables:
            parts.append("")
            parts.append("--- Tables ---")
            for ti, table in enumerate(doc.tables):
                parts.append(f"Table {ti + 1}:")
                for ri, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append("  | " + " | ".join(cells))
                    if ri > 50:  # limit table rows
                        parts.append(f"  ... ({len(table.rows) - 50} more rows)")
                        break
                parts.append("")
            total_parts += len(doc.tables)

        # Embedded images
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        if image_count > 0:
            parts.append(f"--- Note: {image_count} embedded image(s) detected in document ---")
            parts.append("Screenshots or embedded images within this DOCX cannot be read directly. If images contain critical data (charts, configuration screens, signatures), provide them as separate image files for metadata extraction or include a text description.")

        if total_parts == 0:
            return "[DOCX contained no extractable paragraph text or tables]"
        return "\n".join(parts)

    def _extract_image_metadata(self, path: Path) -> str:
        """
        Extract metadata and a text description note for image evidence.

        For screenshot evidence, the model receives the file metadata plus
        a hint that an image was provided. Full OCR is not applied here so
        the user should supplement with text if the image contains critical
        data the model needs to read.
        """
        from PIL import Image

        img = Image.open(path)
        info = [
            f"Image Evidence: {path.name}",
            f"Format: {img.format} | Size: {img.size[0]}x{img.size[1]}px | Mode: {img.mode}",
            "",
            "[IMAGE EVIDENCE NOTE] This evidence was provided as an image/screenshot.",
            "The model cannot read text embedded in images directly.",
            "If this image contains configuration screens, policy text, or audit data,",
            "please include a text description or transcription alongside the image.",
            "The assessment will proceed based on available textual context.",
        ]
        img.close()
        return "\n".join(info)

    # ------------------------------------------------------------------
    # Result construction
    # ------------------------------------------------------------------

    def _build_result(
        self,
        control: dict,
        result_dict: dict,
        statement_type: str,
        tokens_used: int,
    ) -> AssessmentResult:
        """
        Convert the raw Claude response dict into a fully typed AssessmentResult.

        The draft_finding is only materialised when the verdict is FAIL or
        PARTIAL **and** the model returned a non-null finding dict.  For PASS
        and INSUFFICIENT_EVIDENCE the field is always None regardless of what
        the model returned (guards against model hallucination).
        """
        # Safely coerce verdict — fall back to INSUFFICIENT_EVIDENCE if the
        # model returned an unexpected string.
        raw_verdict = result_dict.get("verdict", "INSUFFICIENT_EVIDENCE")
        try:
            verdict = Verdict(raw_verdict)
        except ValueError:
            verdict = Verdict.INSUFFICIENT_EVIDENCE

        # Safely coerce risk rating
        raw_risk = result_dict.get("risk_rating", "INFORMATIONAL")
        try:
            risk_rating = RiskRating(raw_risk)
        except ValueError:
            risk_rating = RiskRating.INFORMATIONAL

        # Build DraftFinding only for actionable verdicts
        draft_finding: DraftFinding | None = None
        raw_finding = result_dict.get("draft_finding")
        if (
            raw_finding
            and isinstance(raw_finding, dict)
            and verdict in (Verdict.FAIL, Verdict.PARTIAL)
        ):
            draft_finding = DraftFinding(
                title=raw_finding.get("title", ""),
                observation=raw_finding.get("observation", ""),
                criteria=raw_finding.get("criteria", ""),
                risk_impact=raw_finding.get("risk_impact", ""),
                recommendation=raw_finding.get("recommendation", ""),
                management_action=raw_finding.get("management_action", ""),
            )

        return AssessmentResult(
            control_id=control["control_id"],
            control_name=control["control_name"],
            statement_type=StatementType(statement_type),
            verdict=verdict,
            confidence=float(result_dict.get("confidence", 0.0)),
            satisfied_requirements=result_dict.get("satisfied_requirements", []),
            gaps=result_dict.get("gaps", []),
            risk_rating=risk_rating,
            draft_finding=draft_finding,
            recommended_evidence=result_dict.get("recommended_evidence", []),
            remediation_notes=result_dict.get("remediation_notes", ""),
            follow_up_questions=result_dict.get("follow_up_questions", []),
            compliance_status=result_dict.get("compliance_status", ""),
            audit_opinion=result_dict.get("audit_opinion", ""),
            assessment_methodology=result_dict.get("assessment_methodology", ""),
            evidence_inventory=result_dict.get("evidence_inventory", []),
            requirement_assessment=result_dict.get("requirement_assessment", []),
            justification=result_dict.get("justification", ""),
            limitations=result_dict.get("limitations", []),
            assessed_at=datetime.utcnow(),
            tokens_used=tokens_used,
            model_used=SONNET,
        )
